from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


DOCS = Path(__file__).resolve().parents[1] / "docs"
CHECKS = {
    "FSD.docx": ("user guide", "docs/user_guide.docx", "report operasional", "laba rugi", "ekspor csv", "rbac"),
    "TSD.docx": ("user guide", "docs/user_guide.docx", "financial_entries", "payment_method", "src/reportsmodule.tsx", "role_permissions", "api/admin/rbac"),
    "USER_GUIDE.docx": ("panduan pelanggan", "panduan cashier", "panduan manager", "panduan admin", "inventory", "add-on", "report operasional", "catat transaksi", "ekspor csv", "pemecahan masalah", "rbac"),
}


for filename, required_terms in CHECKS.items():
    path = DOCS / filename
    assert path.exists() and path.stat().st_size > 20_000, f"File tidak valid: {filename}"
    with ZipFile(path) as archive:
        assert archive.testzip() is None, f"Arsip DOCX korup: {filename}"

    document = Document(path)
    text = " ".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    ).lower()
    missing = [term for term in required_terms if term not in text]
    assert not missing, f"Teks hilang pada {filename}: {missing}"

    for section in document.sections:
        assert abs(section.page_width.inches - 8.5) < 0.01
        assert abs(section.page_height.inches - 11) < 0.01
        for margin in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin):
            assert abs(margin.inches - 1) < 0.01, f"Margin tidak sesuai: {filename}"

    for table_index, table in enumerate(document.tables):
        properties = table._tbl.tblPr
        width = properties.find(qn("w:tblW"))
        indent = properties.find(qn("w:tblInd"))
        assert width is not None and width.get(qn("w:w")) == "9360", f"Lebar tabel {table_index} salah: {filename}"
        assert indent is not None and indent.get(qn("w:w")) == "120", f"Indent tabel {table_index} salah: {filename}"
        grid_widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
        assert sum(grid_widths) == 9360, f"Grid tabel {table_index} salah: {filename}"

print("DOCX content and geometry validation passed:", ", ".join(CHECKS))
