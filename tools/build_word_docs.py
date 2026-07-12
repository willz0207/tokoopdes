from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
OUTPUT_SUFFIX = os.getenv("DOCX_OUTPUT_SUFFIX", "")
OUTPUTS = {
    "fsd": DOCS_DIR / f"FSD{OUTPUT_SUFFIX}.docx",
    "tsd": DOCS_DIR / f"TSD{OUTPUT_SUFFIX}.docx",
    "user_guide": DOCS_DIR / f"USER_GUIDE{OUTPUT_SUFFIX}.docx",
}

PROJECT_NAME = "Franchise Ordering Platform"
UPDATE_DATE = "2026-07-12"

PRIMARY = "2E74B5"
PRIMARY_DARK = "1F4D78"
TEXT = "1F2937"
MUTED = "667085"
LIGHT_FILL = "F4F7FB"
TABLE_HEADER = "E8EEF5"
ACCENT_FILL = "FFF3E8"
BORDER = "D9E2EC"


def hex_to_rgb(value: str) -> RGBColor:
    value = value.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_layout(table, widths: Sequence[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    if widths:
        grid = table._tbl.tblGrid
        for grid_col in list(grid):
            grid.remove(grid_col)
        for width in widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(round(width * 1440)))
            grid.append(grid_col)
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
                    tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:w"), str(round(width * 1440)))
                    tc_w.set(qn("w:type"), "dxa")


def set_table_borders_and_margins(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def apply_run_font(run, size: int | float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = hex_to_rgb(color)


def set_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_to_rgb(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for style_name, size, color, before, after in (
        ("Heading 1", 16, PRIMARY, 18, 10),
        ("Heading 2", 13, PRIMARY, 14, 7),
        ("Heading 3", 12, PRIMARY_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = hex_to_rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.15

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def new_doc(doc_code: str, doc_title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_styles(doc)
    add_footer(section, doc_code)
    add_masthead(doc, doc_code, doc_title, subtitle)
    return doc


def add_footer(section, doc_code: str) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = f"{PROJECT_NAME} · {doc_code} · Update {UPDATE_DATE}"
    for run in paragraph.runs:
        apply_run_font(run, size=8.5, color=MUTED)


def add_masthead(doc: Document, doc_code: str, title: str, subtitle: str) -> None:
    masthead = doc.add_table(rows=1, cols=2)
    set_table_layout(masthead, widths=(4.15, 2.35))
    set_table_borders_and_margins(masthead)
    for cell in masthead.rows[0].cells:
        set_cell_shading(cell, LIGHT_FILL)

    left, right = masthead.rows[0].cells
    p = left.paragraphs[0]
    p.text = doc_code
    p.runs[0].font.size = Pt(9)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = hex_to_rgb(PRIMARY)
    p2 = left.add_paragraph(PROJECT_NAME)
    apply_run_font(p2.runs[0], size=11, bold=True, color=TEXT)
    set_paragraph_spacing(p, after=2, line=1)
    set_paragraph_spacing(p2, after=0, line=1)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.text = f"Update\n{UPDATE_DATE}"
    for run in rp.runs:
        apply_run_font(run, size=9, bold=True, color=MUTED)
    set_paragraph_spacing(rp, after=0, line=1.1)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title_p, before=18, after=4, line=1.1)
    run = title_p.add_run(title)
    apply_run_font(run, size=22, bold=True, color=PRIMARY_DARK)

    subtitle_p = doc.add_paragraph()
    set_paragraph_spacing(subtitle_p, after=14, line=1.2)
    run = subtitle_p.add_run(subtitle)
    apply_run_font(run, size=11, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, after: int = 6) -> None:
    p = doc.add_paragraph(text)
    set_paragraph_spacing(p, after=after)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.25)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        set_paragraph_spacing(p, after=4, line=1.25)


def add_info_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_layout(table, widths=(6.5,))
    set_table_borders_and_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_FILL)
    p = cell.paragraphs[0]
    p.text = title
    apply_run_font(p.runs[0], size=10.5, bold=True, color=PRIMARY_DARK)
    set_paragraph_spacing(p, after=2, line=1.1)
    bp = cell.add_paragraph(body)
    apply_run_font(bp.runs[0], size=10, color=TEXT)
    set_paragraph_spacing(bp, after=0, line=1.2)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_layout(table, widths=widths)
    set_table_borders_and_margins(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], TABLE_HEADER)
        for p in header_cells[idx].paragraphs:
            for run in p.runs:
                apply_run_font(run, size=9.5, bold=True, color=PRIMARY_DARK)
            set_paragraph_spacing(p, after=0, line=1.1)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx])
            set_cell_border(cells[idx])
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    apply_run_font(run, size=9.5, color=TEXT)
                set_paragraph_spacing(p, after=0, line=1.15)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_layout(table, widths=(6.5,))
    set_table_borders_and_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break(WD_BREAK.LINE)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = hex_to_rgb(TEXT)
    doc.add_paragraph()


def build_fsd() -> None:
    doc = new_doc(
        "FSD",
        "Functional Specification Document",
        "Dokumen kebutuhan fungsional untuk platform pemesanan online franchise makanan/minuman yang dapat dikustom.",
    )

    add_heading(doc, "1. Tujuan", 1)
    add_body(
        doc,
        "Aplikasi ini dirancang sebagai platform pemesanan online yang dapat dipakai oleh berbagai franchise makanan dan minuman. "
        "Identitas toko tidak dikunci pada satu brand. Manager dapat menyesuaikan identitas usaha, konten landing page, warna, "
        "kontak, gambar, dan prefix nomor pesanan tanpa perlu mengubah kode aplikasi.",
    )

    add_info_box(
        doc,
        "Prinsip utama",
        "Satu aplikasi dapat digunakan oleh banyak brand atau franchise dengan menyesuaikan konfigurasi, produk, promosi, dan tampilan publik.",
    )

    add_heading(doc, "2. Role Pengguna", 1)
    add_table(
        doc,
        ("Role", "Kapabilitas utama"),
        (
            ("Pelanggan", "Memilih outlet, membuat akun atau login, memilih menu/add-on, menyelesaikan checkout tunai atau online, serta melacak pesanan dan pembayaran."),
            ("Cashier", "Memproses pesanan dan ringkasan operasional hanya dari outlet tempat akunnya ditugaskan."),
            ("Manager", "Mengelola outlet, kategori, produk, promosi, penempatan cashier, inventory, Report, stasiun Cashier, serta pengaturan franchise."),
            ("Admin", "Masuk melalui /admin, mengakses seluruh modul termasuk Outlet, dan mengatur RBAC untuk setiap role."),
        ),
        widths=(1.45, 5.05),
    )

    add_heading(doc, "3. Pengaturan Franchise", 1)
    add_body(doc, "Manager dapat menyesuaikan identitas brand dan konten toko dari pengaturan franchise tanpa perlu mengubah kode aplikasi.")
    add_table(
        doc,
        ("Area", "Yang dapat dikustom"),
        (
            ("Identitas", "Nama usaha, logo singkat, tagline, warna utama, dan warna aksen."),
            ("Kontak", "Nomor WhatsApp toko dan email kontak."),
            ("Pesanan", "Prefix nomor pesanan dan estimasi delivery."),
            ("Landing page", "Teks hero, gambar hero, teks menu, teks tentang, gambar tentang, lokasi, dan footer."),
            ("Dampak tampilan", "Setelah disimpan, perubahan langsung diterapkan pada toko, login, Cashier, pelacakan pesanan, dan dashboard Manager/Admin."),
        ),
        widths=(1.75, 4.75),
    )

    add_heading(doc, "4. Modul Utama", 1)

    add_heading(doc, "4.1 Landing Page Toko", 2)
    add_bullets(
        doc,
        (
            "Menampilkan brand aktif, hero, benefit, katalog menu, promo, tentang, lokasi, dan footer.",
            "Katalog mengambil produk yang aktif, ditugaskan, dan tersedia pada outlet yang sedang dipilih pelanggan.",
            "Filter kategori menu mengambil kategori aktif dari backend dan mengikuti urutan yang diatur Manager.",
            "Jika backend belum aktif, frontend memakai katalog cadangan generik.",
        ),
    )

    add_heading(doc, "4.2 Kategori Menu", 2)
    add_bullets(
        doc,
        (
            "Role yang diberi permission kategori dapat membuat kategori menu custom untuk franchise apa saja.",
            "Setiap kategori memiliki nama, emoji/icon, urutan tampil, dan status aktif.",
            "Dashboard menampilkan kategori dalam kartu yang ringan dan responsif, lengkap dengan ikon, urutan, jumlah produk terkait, serta status visibilitas.",
            "Area tindakan memakai permukaan berwarna lembut, toggle status, dan tombol edit/hapus yang mudah dikenali tanpa membuat kartu terasa berat.",
            "Kategori aktif tampil sebagai filter storefront; kategori nonaktif tidak tampil dan produk pada kategori tersebut tidak dikirim ke katalog publik.",
            "Kategori yang masih dipakai produk akan diarsipkan saat dihapus agar data produk tetap aman.",
            "Rename kategori ikut memperbarui produk terkait secara otomatis.",
        ),
    )

    add_heading(doc, "4.3 Keranjang dan Checkout", 2)
    add_bullets(
        doc,
        (
            "Pelanggan dapat menambah, mengurangi, atau menghapus item dari keranjang.",
            "Produk dapat memiliki add-on opsional; harga add-on dihitung ke harga satuan, subtotal, dan total.",
            "Nama serta harga add-on disimpan sebagai snapshot pada detail pesanan.",
            "Ongkir otomatis gratis jika subtotal minimal Rp75.000.",
            "Pelanggan memilih pembayaran tunai, QRIS, e-wallet, atau transfer bank serta dapat memakai kode promo aktif.",
            "Pelanggan wajib memilih outlet aktif; mengganti outlet akan mengosongkan keranjang agar stok tidak tercampur antar cabang.",
            "Checkout wajib login sebagai pelanggan.",
            "Checkout memvalidasi status master produk, kategori, assignment outlet, status konfigurasi outlet, dan ketersediaan produk.",
            "Harga checkout memakai harga khusus outlet jika diisi; jika tidak, sistem memakai harga master produk.",
            "Pesanan tersimpan di database.",
            "Pembayaran QRIS, e-wallet, dan transfer bank membuat sesi Midtrans; tanpa API key, aplikasi memakai simulator lokal. Sesi pembayaran online Midtrans diatur kedaluwarsa otomatis dalam 15 menit melalui parameter Snap expiry untuk mencegah penguncian stok jangka panjang.",
            "Pesanan online hanya dapat diproses setelah berstatus paid. Pembayaran gagal atau kedaluwarsa membatalkan order dan mengembalikan stok.",
            "Jika nomor WhatsApp toko tersedia, aplikasi membuka WhatsApp dengan pesan konfirmasi.",
        ),
    )

    add_heading(doc, "4.4 Tracking Pesanan", 2)
    add_bullets(
        doc,
        (
            "Pelanggan melihat pesanan aktif dan riwayat pesanan.",
            "Status diperbarui otomatis setiap 15 detik.",
        ),
    )

    add_heading(doc, "4.5 Cashier", 2)
    add_bullets(
        doc,
        (
            "Cashier, Manager, dan Admin dapat memproses pesanan.",
            "Dashboard menampilkan filter order aktif, baru, diproses, siap, dan semua.",
            "Detail pesanan menampilkan add-on yang dipilih pelanggan.",
            "Cashier terkunci pada outlet penempatannya; Manager/Admin dapat berpindah outlet dari selector dashboard.",
        ),
    )

    add_heading(doc, "4.6 Multi-outlet", 2)
    add_bullets(doc, (
        "Manager dan Admin dapat CRUD outlet, menentukan outlet utama, serta memilih konteks cabang aktif.",
        "Outlet menyimpan kode unik, nama, alamat, telepon, status operasional, dan penanda outlet utama.",
        "Pesanan, cashier, inventory, mutasi, transaksi keuangan, dan laporan dipisahkan per outlet.",
        "Kategori, master produk, add-on, promosi, dan brand digunakan bersama. Manager/Admin mengatur assignment, status aktif/tersedia, serta harga khusus produk pada setiap outlet.",
        "Outlet baru menerima assignment awal seluruh produk agar katalog tidak kosong; assignment dapat diubah dari tab Produk.",
        "Outlet utama tidak dapat dihapus. Outlet yang sudah memiliki histori akan diarsipkan agar data tetap utuh.",
    ))

    add_heading(doc, "4.7 Pembayaran Online", 2)
    add_bullets(doc, (
        "Integrasi utama memakai Midtrans Snap sandbox/production dari backend.",
        "Server Key tidak pernah dikirim ke frontend; webhook memverifikasi signature SHA-512 dan nominal order.",
        "Status pembayaran meliputi unpaid, pending, paid, failed, expired, dan refunded.",
        "Simulator lokal tersedia di /payment-simulator saat MIDTRANS_SERVER_KEY belum diisi.",
    ))

    add_heading(doc, "4.8 Inventory", 2)
    add_bullets(
        doc,
        (
            "Manager dan Admin dapat CRUD item inventory per outlet dengan SKU, satuan, stok, batas minimum, dan harga modal. SKU unik di dalam satu outlet.",
            "Item dapat ditautkan ke produk dan jumlah pemakaian agar checkout mengurangi stok secara otomatis.",
            "Pembatalan pesanan (status berubah menjadi cancelled) secara otomatis mengembalikan (refund) stok item terkait ke inventory dan mencatat mutasi masuk (in). Pemulihan pesanan dari batal ke aktif akan memotong stok kembali (gagal jika stok kurang).",
            "Pesanan ditolak secara atomik jika stok item yang tertaut tidak mencukupi.",
            "Mutasi meliputi stok masuk, stok keluar, koreksi tambah, dan koreksi kurang.",
            "Sistem menolak mutasi yang membuat stok negatif dan menyimpan riwayat stok sebelum/sesudah.",
            "Dashboard menampilkan item aktif, stok menipis, mutasi hari ini, dan riwayat mutasi.",
        ),
    )

    add_heading(doc, "4.9 Report Operasional dan Keuangan", 2)
    add_bullets(
        doc,
        (
            "Report operasional mencakup penjualan harian/real-time, produk terlaris, pembayaran, stok, dan pelanggan.",
            "Report keuangan mencakup laba rugi (menampilkan Pendapatan, HPP/COGS otomatis berdasarkan resep/modal bahan, Laba Kotor, Biaya operasional manual, dan Laba Bersih), arus kas, neraca sederhana, perubahan modal, dan biaya per kategori.",
            "Manager/Admin dapat mencatat biaya atau perubahan modal, memfilter periode, dan mengekspor CSV.",
            "Penjualan tunai non-batal dan pembayaran online berstatus paid menjadi pendapatan otomatis serta menghitung HPP. Report selalu mengikuti outlet aktif.",
        ),
    )

    add_heading(doc, "4.10 Manager dan Admin", 2)
    add_bullets(
        doc,
        (
            "Seluruh modul dashboard Manager/Admin memakai pola visual yang sama: panel terang, sudut membulat, jarak antarelemen yang lega, header dan area tindakan yang konsisten, serta state hover/focus yang jelas.",
            "Tampilan tetap nyaman dipakai pada desktop maupun layar ponsel. Navigasi berubah menjadi baris menu bawah dan kartu/tabel menyesuaikan lebar layar tanpa menimbulkan scroll horizontal halaman.",
            "Tab Kategori untuk CRUD kategori menu jika role memiliki permission kategori.",
            "Tab Produk untuk CRUD master produk, pengelolaan add-on, dan assignment produk pada outlet aktif.",
            "Tab Promosi untuk CRUD promosi.",
            "Tab Cashier untuk CRUD akun cashier, termasuk edit identitas/password dan pengaturan status akses login.",
            "Daftar cashier memakai kartu responsif dengan avatar, nama, email, status akses login, footer terang, serta tombol Edit dan Hapus yang mudah dikenali.",
            "Tab Inventory untuk item stok, batas minimum, dan stock movement.",
            "Tab Report untuk laporan operasional/keuangan, ekspor CSV, serta transaksi biaya/modal.",
            "Tab Outlet untuk CRUD cabang dan memilih konteks Cashier, Inventory, serta Report.",
            "Tab Franchise untuk pengaturan brand dan halaman publik.",
            "Tab RBAC khusus Admin untuk mengatur akses modul per role.",
        ),
    )

    add_heading(doc, "4.11 RBAC Modul", 2)
    add_bullets(
        doc,
        (
            "Admin dapat membuka tab RBAC pada dashboard Manager/Admin.",
            "RBAC menyimpan matrix akses role cashier, manager, dan admin terhadap modul Stasiun Cashier, Kategori Menu, Produk & Add-on, Promosi, Akun Cashier, Inventory, Report, Outlet, Franchise, dan RBAC.",
            "Menu dashboard hanya menampilkan tab yang diizinkan untuk role aktif.",
            "Backend memvalidasi permission modul pada setiap endpoint operasional.",
            "Permission Admin - RBAC wajib aktif agar Admin tidak terkunci dari pengaturan akses.",
            "Permission RBAC untuk Cashier/Manager dikunci nonaktif; pengaturan RBAC hanya dapat dikelola Admin.",
            "Tabel akses memakai baris yang lega, penanda aktif/nonaktif yang mudah dibaca, serta panel tindakan terang dan ringkas. Panel tersebut memuat pengingat singkat, tombol Reset tampilan, dan tombol Simpan RBAC tanpa menutupi isi halaman.",
        ),
    )

    add_heading(doc, "4.12 Keamanan Sistem", 2)
    add_bullets(
        doc,
        (
            "Verifikasi Kekuatan Kata Sandi: Pengguna (pelanggan maupun kasir) wajib menggunakan kata sandi minimal 8 karakter dengan kombinasi huruf besar, huruf kecil, angka, dan simbol khusus.",
            "Pembatasan Percobaan Login (Rate Limiting): Setiap IP dibatasi maksimal 5 kali request login/registrasi per 15 menit. Jika dilanggar, server mengembalikan respons HTTP 429.",
            "Kebijakan Keamanan Konten (CSP): Mengaktifkan CSP melalui Helmet untuk mencegah eksekusi script luar yang tidak sah, dengan tetap mengizinkan fungsi pembayaran online Midtrans Snap.",
        ),
    )

    add_heading(doc, "4.12 Mode Development Lokal", 2)
    add_bullets(
        doc,
        (
            "Aplikasi dapat berjalan lokal tanpa Netlify.",
            "npm run dev menjalankan API Express lokal pada http://localhost:3001 dan frontend Vite pada http://localhost:5175.",
            "Route frontend /api/* diproxy ke API lokal sehingga pengguna tetap membuka satu alamat utama: http://localhost:5175.",
            "Database development memakai PostgreSQL lokal dari DATABASE_URL, POSTGRES_URL, atau LOCAL_DATABASE_URL.",
            "Jika tidak ada connection string, backend memakai default postgres://postgres@127.0.0.1:5432/postgres.",
            "Migration awal tetap memakai file SQL project dan diterapkan otomatis saat database kosong.",
        ),
    )

    add_heading(doc, "4.13 Hosting Publik Opsional", 2)
    add_bullets(
        doc,
        (
            "Project masih dapat dibuild untuk hosting publik melalui Netlify Function atau host Node lain.",
            "Database lokal di komputer tidak dapat dipakai oleh website publik ketika komputer mati/offline atau tidak dapat diakses dari internet.",
            "Untuk website publik yang aktif 24 jam, hosting perlu memakai PostgreSQL online yang diisi melalui DATABASE_URL atau POSTGRES_URL.",
            "Password Admin, Manager, Cashier, dan JWT secret harus disimpan sebagai environment variable hosting dan tidak masuk repository.",
        ),
    )

    add_heading(doc, "5. Aturan Dokumen", 1)
    add_body(doc, "Setiap perubahan fitur, alur bisnis, API, database, role, atau UI utama harus memperbarui dokumen berikut:")
    add_bullets(doc, ("docs/FSD.md", "docs/TSD.md", "docs/USER_GUIDE.md", "docs/FSD.docx", "docs/TSD.docx", "docs/USER_GUIDE.docx"))

    add_heading(doc, "6. Riwayat Perubahan", 1)
    add_table(
        doc,
        ("Tanggal", "Perubahan"),
        (
            ("2026-07-12", "Mengganti pemilih outlet bawaan browser dengan komponen dropdown kustom (OutletSelector) yang lebih interaktif dan premium di storefront, manager, dan cashier."),
            ("2026-07-12", "Meningkatkan keamanan web: menambahkan in-memory rate limiter untuk login/register, menerapkan validasi kekuatan kata sandi (huruf besar/kecil, angka, simbol), dan mengonfigurasi Content Security Policy (CSP) pada Helmet."),
            ("2026-07-12", "Mengoptimasi performa backend: meningkatkan pg pool size ke 20, mengatasi bottleneck kueri N+1 pada produk via bulk fetch addons, mengimplementasi cache matriks otorisasi role, menggabungkan kueri statistik dashboard dengan CTE, mempercepat cold-start migrasi via check to_regclass, serta menambahkan masa berlaku pembayaran Snap 15 menit."),
            ("2026-07-12", "Menambahkan assignment produk per outlet, harga khusus, status aktif/tersedia per cabang, filter katalog outlet, dan validasi checkout berbasis outlet_products, sementara kategori dan master produk tetap global."),
            ("2026-07-12", "Menyeragamkan tampilan seluruh modul Manager/Admin—Produk, Kategori, Promosi, Cashier, Inventory, Report, Outlet, Franchise, dan RBAC—dengan panel terang, area tindakan konsisten, serta layout responsif."),
            ("2026-07-12", "Menambahkan multi-outlet untuk pesanan, cashier, inventory, transaksi keuangan, dan laporan; melengkapi dashboard dengan pemilih outlet serta modul CRUD cabang."),
            ("2026-07-12", "Menambahkan pembayaran online Midtrans Snap, verifikasi webhook, status pembayaran, perlindungan proses pesanan sebelum lunas, dan simulator lokal tanpa API key."),
            ("2026-07-12", "Menyegarkan tampilan RBAC dengan tabel yang lebih nyaman dibaca, switch yang lebih jelas, dan panel simpan terang yang ringkas serta responsif."),
            ("2026-07-12", "Menyelaraskan tampilan kartu cashier dengan kartu kategori: avatar lembut, status akses informatif, footer terang, tombol aksi berlabel, serta efek hover."),
            ("2026-07-12", "Menyegarkan tampilan kartu kategori dengan layout yang lebih ringan, status visibilitas informatif, footer lembut, tombol aksi interaktif, dan efek hover."),
            ("2026-07-12", "Merapikan gaya bahasa FSD agar lebih natural dan mudah dibaca tanpa mengubah makna fungsional maupun istilah teknis."),
            ("2026-07-12", "Memisahkan URL dashboard berdasarkan role: Admin memakai /admin, Manager memakai /manager, serta menambahkan pengalihan otomatis ke URL kanonis role."),
            ("2026-07-10", "Mengubah mode development agar database berjalan lokal tanpa Netlify: API lokal port 3001, Vite port 5175 dengan proxy /api, PostgreSQL lokal melalui DATABASE_URL, dan dokumentasi lokal diperbarui."),
            ("2026-07-06", "Deploy production Netlify aktif pada https://tokokopdes.netlify.app, Netlify Database diprovision, migration production diterapkan, dan backend dibuat kompatibel dengan deploy manual CLI melalui NETLIFY_DB_URL."),
            ("2026-07-06", "Memigrasikan hosting ke Netlify Free, Express ke Netlify Functions, SQLite ke Netlify Database/PostgreSQL persisten, migration otomatis, routing SPA/API, dan emulator lokal Vite."),
            ("2026-07-06", "Menambahkan Blueprint Render untuk hosting publik gratis, domain onrender.com, secret credential saat deploy, health check, auto-deploy, dan batasan penyimpanan SQLite sementara."),
            ("2026-07-05", "Menambahkan RBAC per modul untuk Admin, tabel permission role, tab RBAC, menu dashboard berbasis permission, dan validasi permission di backend."),
            ("2026-07-05", "Menambahkan modul kategori menu custom berbasis database, endpoint kategori publik/manager, tab Kategori khusus Manager, dan sinkronisasi kategori ke storefront."),
            ("2026-07-05", "Menambahkan modul Report, pembayaran/promo checkout, transaksi biaya/modal, ekspor CSV, valuasi persediaan, dan pengurangan stok otomatis."),
            ("2026-07-05", "Menambahkan User Guide sebagai dokumen wajib yang selalu diperbarui bersama FSD dan TSD."),
            ("2026-07-05", "Menambahkan add-on produk end-to-end, role Admin dengan akses seluruh modul operasional, dan modul Inventory untuk item stok serta stock movement."),
            ("2026-07-05", "Melengkapi modul Cashier pada role Manager menjadi CRUD penuh: tambah, lihat, edit, aktif/nonaktif, ubah password, dan hapus."),
            ("2026-07-05", "Menambahkan penyempurnaan stabilitas agar setup database baru tetap dapat menjalankan modul manager dan promosi."),
            ("2026-07-05", "Menambahkan modul manager untuk melihat daftar cashier dan membuat akun cashier baru."),
            ("2026-07-05", "Project digenerikkan untuk franchise apa saja, ditambah pengaturan franchise, upload gambar brand, dan dokumen FSD/TSD."),
            ("2026-07-05", "Sisa teks/konfigurasi brand lama dibersihkan dari UI, README, env example, seed/default, dan key penyimpanan browser."),
            ("2026-07-05", "Dokumen FSD/TSD dibuat ulang dalam format Word yang lebih rapi untuk kebutuhan review dan arsip."),
        ),
        widths=(1.35, 5.15),
    )

    doc.save(OUTPUTS["fsd"])


def build_tsd() -> None:
    doc = new_doc(
        "TSD",
        "Technical Specification Document",
        "Dokumen spesifikasi teknis untuk frontend, backend, database, API, role, dan konfigurasi franchise.",
    )

    add_heading(doc, "1. Stack Teknologi", 1)
    add_body(doc, "Bagian ini merangkum susunan teknis aplikasi, mulai dari frontend dan backend hingga database, autentikasi, serta pola deployment yang didukung.")
    add_table(
        doc,
        ("Layer", "Teknologi"),
        (
            ("Frontend", "Antarmuka dibangun dengan React, TypeScript, dan Vite."),
            ("Backend", "Express dan TypeScript; development lokal berjalan melalui server/start.ts, sedangkan Netlify Function tetap tersedia sebagai opsi hosting."),
            ("Database", "PostgreSQL melalui package pg; development memakai database lokal tanpa bergantung pada Netlify."),
            ("Auth", "Autentikasi berbasis JWT."),
            ("Styling", "Setiap halaman utama memiliki file CSS tersendiri agar styling lebih mudah dirawat."),
        ),
        widths=(1.65, 4.85),
    )

    add_heading(doc, "2. Struktur Penting", 1)
    add_table(
        doc,
        ("Path", "Fungsi"),
        (
            ("src/App.tsx", "Landing page, pemilih outlet, katalog, keranjang, checkout, dan redirect pembayaran online."),
            ("src/AuthApp.tsx", "Login/register pelanggan serta login cashier, manager, admin."),
            ("src/roleRoutes.ts", "Pemetaan URL halaman utama berdasarkan role pengguna."),
            ("src/CashierApp.tsx", "Dashboard cashier."),
            ("src/CustomerOrdersApp.tsx", "Tracking pesanan pelanggan."),
            ("src/ManagerApp.tsx", "Dashboard manager/admin berbasis permission, pemilih/CRUD outlet, master produk, assignment produk per outlet, kategori, promosi, cashier, inventory, Report, RBAC, dan settings."),
            ("src/PaymentSimulatorApp.tsx", "Simulator pembayaran lokal tanpa API key Midtrans."),
            ("src/manager.css", "Sistem visual bersama untuk seluruh dashboard Manager/Admin: token surface/shadow, panel, kartu, tabel, toolbar, modal, action area, hover/focus, dan layout responsif."),
            ("src/InventoryModule.tsx", "UI item inventory, minimum stok, mutasi, dan riwayat stok."),
            ("src/ReportsModule.tsx", "Dashboard Report operasional/keuangan, filter periode, ekspor CSV, dan transaksi biaya/modal."),
            ("src/franchise.ts", "Default settings, hook settings, dan apply warna brand."),
            ("src/api.ts", "API client frontend."),
            ("server/index.ts", "Routing API async, middleware auth/RBAC, dan Express app untuk Function."),
            ("server/start.ts", "Entry point API lokal pada port 3001 untuk npm run dev."),
            ("server/contracts.ts", "Kontrak tipe domain bersama untuk API dan provider database."),
            ("server/postgres-db.ts", "Query PostgreSQL, transaksi, seed idempotent, dan agregasi laporan."),
            ("server/payments.ts", "Adapter Midtrans Snap, simulator lokal, pemetaan status, dan verifikasi webhook."),
            ("netlify/functions/api.ts", "Entry point Express Netlify Function."),
            ("netlify/database/migrations/0001_initial_schema.sql", "Baseline schema PostgreSQL yang diterapkan otomatis."),
            ("netlify/database/migrations/0002_multi_outlet_payments.sql", "Migrasi outlet, scoping data operasional, penempatan cashier, dan pembayaran."),
            ("netlify/database/migrations/0003_outlet_products.sql", "Migrasi assignment produk per outlet, harga override, status aktif, dan status tersedia."),
            ("netlify.toml", "Build, bundling Function, routing API, dan fallback SPA."),
            ("vite.config.ts", "Konfigurasi Vite, port 5175, dan proxy /api ke API lokal port 3001."),
        ),
        widths=(2.25, 4.25),
    )

    add_heading(doc, "3. Database", 1)
    add_body(doc, "Pada mode development, aplikasi terhubung langsung ke PostgreSQL lokal tanpa melalui Netlify. Provider database mencari connection string dari DATABASE_URL, POSTGRES_URL, lalu LOCAL_DATABASE_URL. Jika ketiganya kosong, aplikasi memakai fallback postgres://postgres@127.0.0.1:5432/postgres. Variable NETLIFY_DB_URL tidak lagi digunakan.")
    add_body(doc, "Perintah npm run dev menjalankan dua proses lokal secara bersamaan:")
    add_bullets(
        doc,
        (
            "npm run dev:api: tsx watch server/start.ts untuk API Express lokal pada port 3001.",
            "npm run dev:web: Vite pada port 5175, dengan proxy /api ke http://127.0.0.1:3001.",
        ),
    )
    add_table(
        doc,
        ("Tabel", "Kegunaan"),
        (
            ("users", "Data akun pelanggan, cashier, manager, dan admin."),
            ("menu_categories", "Kategori menu custom, emoji/icon, urutan tampil, status aktif, dan jumlah produk terkait."),
            ("products", "Data menu/product yang tampil di katalog."),
            ("product_addons", "Pilihan add-on, harga, status, dan relasi ke produk."),
            ("outlet_products", "Assignment produk, harga khusus, status aktif, dan ketersediaan per outlet."),
            ("promotions", "Data promo yang dapat dikelola manager."),
            ("orders", "Header pesanan, status, total, dan data pengiriman."),
            ("order_items", "Detail produk, harga satuan, dan snapshot add-on per pesanan."),
            ("outlets", "Kode, nama, alamat, telepon, status, dan outlet utama."),
            ("payments", "Provider, status, token, redirect, transaksi, nominal, dan raw response pembayaran."),
            ("inventory_items", "Outlet, SKU unik per cabang, stok, minimum, harga modal, relasi produk, dan status."),
            ("stock_movements", "Mutasi stok beserta stok sebelum/sesudah dan pembuat."),
            ("financial_entries", "Biaya operasional, modal masuk/keluar, kategori, metode pembayaran, dan tanggal."),
            ("role_permissions", "Matrix permission role terhadap modul operasional dan RBAC."),
            ("app_settings", "Pengaturan franchise dalam bentuk key-value."),
        ),
        widths=(1.8, 4.7),
    )

    add_heading(doc, "3.1 RBAC", 2)
    add_bullets(
        doc,
        (
            "role_permissions menyimpan role, module, enabled, dan updated_at.",
            "Role yang dapat diatur adalah cashier, manager, dan admin.",
            "Modul RBAC: cashier_station, categories, products, promotions, cashiers, inventory, reports, outlets, settings, dan rbac.",
            "Default permission: Cashier hanya cashier_station; Manager seluruh modul operasional kecuali rbac; Admin seluruh modul termasuk rbac.",
            "Permission admin + rbac dipaksa aktif; permission rbac untuk Cashier/Manager dipaksa nonaktif.",
            "server/index.ts memakai requireModuleAccess dan requireAnyModuleAccess agar permission ditegakkan di backend.",
            "GET /api/permissions/me dipakai frontend untuk menentukan tab dashboard yang boleh tampil.",
            "ManagerApp.tsx merender panel tindakan sebagai div.rbac-actions, bukan footer, agar tidak menerima gaya footer global dari bagian lain aplikasi.",
            "manager.css memberi tabel, switch, dan panel tindakan RBAC state hover, focus-visible, serta susunan responsif. Input switch tetap tersedia bagi keyboard dan pembaca layar.",
        ),
    )

    add_heading(doc, "3.2 Sistem Visual Dashboard Manager/Admin", 2)
    add_bullets(
        doc,
        (
            "Selector footer halaman publik dibatasi menjadi .app-shell > footer di src/styles.css agar warna footer storefront tidak diterapkan ke footer kartu atau panel dashboard.",
            "src/manager.css menyediakan token bersama --module-surface, --module-soft, --module-border, dan beberapa tingkat shadow agar warna, garis tepi, kedalaman, dan radius setiap modul tetap seragam.",
            "Panel, ringkasan statistik, toolbar, tabel, modal, empty state, serta area tindakan pada Produk, Promosi, Inventory, Report, Outlet, Franchise, dan RBAC memakai pola interaksi yang konsisten, termasuk hover, focus-visible, dan disabled state.",
            "Media query dashboard mengubah grid menjadi satu kolom dan tindakan menjadi susunan vertikal pada layar kecil. Pengujian viewport 390 px memastikan halaman tidak menghasilkan overflow horizontal.",
        ),
    )

    add_heading(doc, "3.3 Kategori Menu", 2)
    add_bullets(
        doc,
        (
            "menu_categories menyimpan label, emoji, sort_order, status aktif, created_at, dan updated_at.",
            "Seed kategori default dibuat saat database kosong; kategori unik dari produk lama ikut dimigrasikan.",
            "GET /api/categories?outletId=:id hanya mengirim kategori aktif yang memiliki produk aktif dan tersedia pada outlet terpilih.",
            "GET /api/manager/categories mengirim kategori aktif/nonaktif beserta productCount.",
            "CategoryGrid di src/ManagerApp.tsx memisahkan informasi utama, status visibilitas, toggle aktif, dan tombol aksi ke dalam kelompok visual yang mudah dipindai.",
            "src/manager.css memberi kartu kategori layout responsif tiga/dua/satu kolom, warna status aktif/nonaktif, footer terang, dan transisi hover.",
            "CRUD kategori dibatasi permission categories melalui POST/PUT/PATCH/DELETE /api/manager/categories.",
            "Produk divalidasi terhadap menu_categories dan tidak lagi memakai daftar kategori hardcoded.",
            "getProducts(false, ..., outletId) hanya mengirim produk yang master dan kategorinya aktif serta memiliki assignment aktif/tersedia; getProducts(true, ..., outletId) menampilkan seluruh master produk beserta assignment outlet.",
            "Rename kategori ikut memperbarui products.category; delete kategori terpakai mengarsipkan kategori dengan active = 0.",
        ),
    )

    add_heading(doc, "3.4 Add-on dan Detail Pesanan", 2)
    add_bullets(
        doc,
        (
            "API publik hanya mengirim add-on aktif; Manager/Admin menerima add-on aktif dan nonaktif.",
            "order_items.addons_json menyimpan snapshot ID, nama, dan harga add-on saat checkout.",
            "order_items.unit_price menyimpan harga produk ditambah add-on per satuan.",
        ),
    )

    add_heading(doc, "3.5 Produk per Outlet", 2)
    add_bullets(
        doc,
        (
            "outlet_products memakai primary key outlet_id + product_id serta menyimpan price_override, active, available, dan timestamp.",
            "Migrasi 0003_outlet_products.sql menugaskan seluruh produk lama ke seluruh outlet agar perilaku katalog lama tetap terjaga.",
            "Outlet baru menerima assignment seluruh master produk; produk baru otomatis ditugaskan ke outlet aktif saat dibuat.",
            "GET /api/products?outletId=:id mengembalikan harga efektif COALESCE(price_override, products.price) dan hanya produk yang tersedia pada outlet.",
            "GET /api/manager/products memakai X-Outlet-Id dan mengembalikan basePrice serta outletAssignment.",
            "PUT /api/manager/products/:id/outlet-assignment membuat, memperbarui, atau menghapus assignment outlet.",
            "createOrder mengambil ulang produk berdasarkan outlet, memvalidasi assignment, dan menghitung subtotal memakai harga efektif server-side.",
        ),
    )

    add_heading(doc, "3.6 Inventory", 2)
    add_bullets(
        doc,
        (
            "Tipe mutasi: in, out, adjustment_add, adjustment_subtract.",
            "inventory_items memiliki outlet_id dan unique index outlet_id + UPPER(sku).",
            "Update current_stock dan insert stock_movements berjalan dalam satu transaksi PostgreSQL dengan row locking.",
            "Mutasi ditolak jika menghasilkan stok negatif.",
            "Order memvalidasi master produk, kategori, assignment outlet, status aktif/tersedia, dan harga efektif. Item terkait dikurangi sebesar quantity x usage_per_sale dan stock movement dibuat otomatis.",
            "Ketika pesanan dibatalkan (status berubah menjadi cancelled), stok dikembalikan (refund) secara otomatis dan dicatat sebagai mutasi masuk (in). Sebaliknya, jika pesanan dipulihkan kembali dari status batal, stok akan dikurangi kembali (gagal jika stok tidak mencukupi).",
        ),
    )

    add_heading(doc, "3.7 Report dan Keuangan", 2)
    add_bullets(
        doc,
        (
            "orders menyimpan payment_method, discount_amount, dan promo_code sebagai snapshot checkout.",
            "Agregasi laporan dikerjakan server-side berdasarkan periode dan outlet; order cancelled serta pembayaran online belum lunas dikecualikan.",
            "Laba rugi: omzet bersih dikurangi HPP (COGS) dan biaya operasional periode untuk menghasilkan laba bersih. HPP dihitung dinamis dari total penjualan produk dikali resep bahan baku dikali unit_cost inventory item.",
            "Laba kotor: omzet bersih dikurangi HPP (COGS).",
            "Neraca sederhana memakai saldo kas historis dan nilai inventory current_stock x unit_cost.",
        ),
    )

    add_heading(doc, "3.8 app_settings", 2)
    add_body(doc, "Tabel app_settings menyimpan pengaturan franchise dalam bentuk key-value. Field yang dipakai aplikasi:")
    add_bullets(
        doc,
        (
            "businessName, shortName, tagline, heroEyebrow, heroTitle, heroHighlight, heroDescription",
            "heroImageUrl, storyImageUrl, deliveryEstimate, deliveryNote",
            "locationLabel, locationTitle, locationDescription, footerDescription",
            "contactEmail, whatsappNumber, orderPrefix, primaryColor, accentColor",
            "menuKicker, menuTitle, menuDescription, aboutKicker, aboutTitle",
            "aboutDescription, aboutReviewQuote, aboutReviewAuthor",
        ),
    )

    add_heading(doc, "3.9 Multi-outlet", 2)
    add_bullets(doc, (
        "outlets menyimpan code, name, address, phone, active, is_default, dan timestamp.",
        "users.outlet_id menempatkan Cashier/Manager; orders, inventory_items, dan financial_entries memiliki outlet_id.",
        "Cashier selalu memakai outlet akun. Manager/Admin berpindah konteks lewat header X-Outlet-Id.",
        "Data lama dimigrasikan ke Outlet Pusat. Kategori, master produk, add-on, promosi, RBAC, dan app_settings tetap global; ketersediaan serta harga produk dipisahkan melalui outlet_products.",
    ))

    add_heading(doc, "3.10 Pembayaran Online", 2)
    add_bullets(doc, (
        "payments berelasi unik ke orders dan menyimpan provider, status, token, redirect URL, transaction ID, nominal, raw response, dan paid_at.",
        "server/payments.ts memakai Midtrans sandbox/production atau simulator saat MIDTRANS_SERVER_KEY kosong.",
        "Webhook memvalidasi SHA-512 order_id + status_code + gross_amount + server_key serta mencocokkan nominal order.",
        "Status gagal/kedaluwarsa membatalkan order dan mengembalikan stok; order online belum paid tidak dapat diproses. Sesi pembayaran online Midtrans diatur kedaluwarsa otomatis dalam 15 menit melalui parameter Snap expiry untuk mencegah penguncian stok jangka panjang.",
    ))

    add_heading(doc, "3.11 Optimasi Performa Database dan Serverless", 2)
    add_bullets(doc, (
        "Eliminasi Kueri N+1 pada Produk: Fungsi getProducts mem-bulk fetch add-on untuk semua produk dalam satu kueri SQL menggunakan ANY($1::int[]) dan memetakan di memori.",
        "Caching Otorisasi Dinamis: Hak akses disimpan di memori (rolePermissionsCache) dan dibersihkan otomatis ketika matriks RBAC diperbarui.",
        "Peningkatan Koneksi Pool: Connection pool PostgreSQL dinaikkan menjadi max: 20 untuk mendukung beban kueri analitik dashboard.",
        "Optimasi Kueri Laporan: Kueri statistik dashboard digabung dengan CTE, dan pengeluaran dikalkulasi di memori dari entri transaksi.",
        "Verifikasi Migrasi Cepat: seedDatabase menggunakan to_regclass untuk mendeteksi keberadaan tabel, mempercepat cold start serverless.",
    ))

    add_heading(doc, "3.12 Fitur Keamanan Tambahan", 2)
    add_bullets(doc, (
        "Pembatasan Request (Rate Limiting): Endpoint autentikasi dilindungi oleh middleware rateLimiter in-memory (maksimal 5 kali percobaan per 15 menit per IP).",
        "Kebijakan Kekuatan Kata Sandi: Setiap pembuatan/pembaruan kata sandi divalidasi oleh isStrongPassword (minimal 8 karakter, ada huruf besar, huruf kecil, angka, dan simbol khusus).",
        "Content Security Policy (CSP): Helmet dikonfigurasi dengan direktif CSP yang aman namun tetap kompatibel dengan Vite development server dan Midtrans Snap.",
    ))

    add_heading(doc, "4. API Utama", 1)
    add_table(
        doc,
        ("Area", "Endpoint"),
        (
            ("Public", "GET /api/health; GET /api/outlets; GET settings/categories/products/promotions; POST auth register/login"),
            ("Customer", "POST /api/orders; GET customer orders/payment status; POST simulator payment; profile endpoints"),
            ("Cashier / Manager / Admin", "GET /api/staff/outlets; GET cashier stats/orders; PATCH order status"),
            ("Manager / Admin", "CRUD outlet/category/product/add-on/promotion/cashier; inventory; reports; financial entries; settings"),
            ("Permission / Admin RBAC", "GET /api/permissions/me; GET /api/admin/rbac; PUT /api/admin/rbac"),
            ("Webhook", "POST /api/payments/midtrans/notification dengan verifikasi signature dan nominal"),
        ),
        widths=(1.75, 4.75),
    )

    add_heading(doc, "5. Auth dan Role", 1)
    add_body(doc, "JWT memakai secret dari APP_JWT_SECRET. Role yang didukung: customer, cashier, manager, dan admin.")
    add_body(doc, "Admin disimpan sebagai user biasa. Setelah login, homePathForRole memilih halaman utama berdasarkan role: Customer ke /, Cashier ke /cashier, Manager ke /manager, dan Admin ke /admin. Rute /admin dan /manager sama-sama menggunakan ManagerApp, tetapi aplikasi tetap mengalihkan pengguna jika URL tidak sesuai dengan role sesi. Endpoint /api/auth/admin dipertahankan untuk kompatibilitas API lama.")
    add_body(doc, "Untuk setiap token yang memiliki userId, middleware memastikan akun masih ada, masih aktif, dan role di database sesuai dengan isi token. Hasil kueri user tersebut disimpan dalam request context (request.user) dan digunakan kembali di resolveOutletId untuk menghemat kueri database sekuensial. Sesi lama tidak dapat digunakan lagi setelah akun dinonaktifkan atau dihapus.")
    add_body(doc, "Frontend menyimpan token dan data sesi di localStorage dengan key generik:")
    add_table(
        doc,
        ("Key", "Fungsi"),
        (
            ("franchise-user-token", "Token pelanggan, cashier, manager, atau admin."),
            ("franchise-user", "Data profil user aktif."),
            ("franchise-admin-token", "Token halaman admin legacy."),
            ("franchise-cart", "Keranjang belanja."),
            ("franchise-outlet-id", "Outlet aktif pada storefront dan dashboard."),
        ),
        widths=(2.35, 4.15),
    )

    add_heading(doc, "5.1 Enforcement RBAC", 2)
    add_bullets(
        doc,
        (
            "Endpoint Cashier membutuhkan permission cashier_station.",
            "Endpoint kategori membutuhkan permission categories; GET /api/manager/categories juga dapat dipakai permission products untuk dropdown produk.",
            "Endpoint produk membutuhkan permission products; GET /api/manager/products juga dapat dipakai permission inventory untuk tautan item inventory.",
            "PUT /api/manager/products/:id/outlet-assignment membutuhkan permission products dan konteks outlet yang valid.",
            "Endpoint promosi, cashier, inventory, report/financial, outlet, dan settings membutuhkan permission modul terkait.",
            "Endpoint RBAC membutuhkan role admin dan permission rbac.",
            "Endpoint legacy /api/admin/* tetap tersedia dan ikut dicek memakai permission modul terkait.",
        ),
    )

    add_heading(doc, "6. Upload Gambar", 1)
    add_body(doc, "Foto produk dan gambar brand dikirim sebagai Data URL base64 agar dapat disimpan bersama data aplikasi tanpa layanan penyimpanan file tambahan.")
    add_bullets(
        doc,
        (
            "Format: PNG, JPG/JPEG, WebP, GIF.",
            "Maksimal sekitar 2 MB per gambar.",
            "Server membatasi JSON request hingga 6 MB.",
            "Cara ini praktis untuk penggunaan lokal atau demo. Untuk produksi dengan banyak gambar, file sebaiknya dipindahkan ke object storage atau folder upload statis.",
        ),
    )

    add_heading(doc, "7. Pengaturan Brand Frontend", 1)
    add_body(doc, "Frontend mengambil settings dari GET /api/settings melalui useFranchiseSettings(). Hook tersebut:")
    add_bullets(
        doc,
        (
            "Menggunakan default generik saat data belum tersedia.",
            "Mengambil settings dari backend.",
            "Mengubah CSS variable --brand-primary dan --brand-accent.",
            "Mengubah judul browser sesuai nama usaha.",
        ),
    )

    add_heading(doc, "8. Build dan Run", 1)
    add_bullets(doc, (
        "MIDTRANS_SERVER_KEY adalah credential rahasia backend; jika kosong aplikasi memakai simulator lokal.",
        "MIDTRANS_CLIENT_KEY disiapkan untuk Snap JS; MIDTRANS_IS_PRODUCTION menentukan sandbox/production.",
        "PUBLIC_APP_URL menjadi base URL callback selesai pembayaran.",
    ))
    add_table(
        doc,
        ("Mode", "Command"),
        (
            ("Development", "npm run dev"),
            ("Build", "npm run build"),
            ("Production", "npm run build; host publik menerbitkan dist dan backend sesuai platform."),
        ),
        widths=(1.65, 4.85),
    )

    add_heading(doc, "8.1 Hosting Publik Opsional", 2)
    add_bullets(
        doc,
        (
            "Production URL sebelumnya: https://tokokopdes.netlify.app.",
            "Infrastructure as Code memakai netlify.toml; build command npm run build dan publish directory dist.",
            "netlify/functions/api.ts membungkus Express dengan serverless-http; /api/* di-rewrite ke Function dan /* memakai fallback SPA.",
            "server/postgres-db.ts memakai connection string eksplisit dari environment (DATABASE_URL, POSTGRES_URL, atau LOCAL_DATABASE_URL) dan tidak lagi memakai NETLIFY_DB_URL.",
            "Seed settings, role permission, akun awal, kategori, produk, dan promosi bersifat idempotent dengan advisory lock.",
            "Checkout dan stock movement memakai transaksi PostgreSQL serta FOR UPDATE untuk mencegah overselling.",
            "Local development tidak memakai @netlify/vite-plugin; API lokal dan Vite lokal dijalankan terpisah oleh concurrently.",
            "Database lokal komputer tidak dapat dipakai oleh website publik jika komputer mati/offline. Untuk hosting publik 24 jam, gunakan PostgreSQL online yang dapat diakses oleh server hosting.",
            "Secret produksi disimpan pada environment variable hosting: APP_JWT_SECRET serta password Admin, Manager, dan Cashier.",
            "Jika tetap memakai Netlify Free, perhatikan batas kredit bulanan karena project dapat pause otomatis jika limit tercapai.",
        ),
    )

    add_heading(doc, "9. Aturan Update Dokumen", 1)
    add_info_box(
        doc,
        "Wajib sinkron",
        "Setiap perubahan kode yang memengaruhi fitur, API, database, role, pengaturan brand, upload, atau alur checkout wajib memperbarui FSD, TSD, dan User Guide, termasuk versi Markdown dan Word.",
    )
    add_bullets(doc, ("docs/FSD.md", "docs/TSD.md", "docs/USER_GUIDE.md", "docs/FSD.docx", "docs/TSD.docx", "docs/USER_GUIDE.docx"))

    add_heading(doc, "10. Riwayat Perubahan", 1)
    add_table(
        doc,
        ("Tanggal", "Perubahan teknis"),
        (
            ("2026-07-12", "Mengganti pemilih outlet bawaan browser dengan komponen dropdown kustom (OutletSelector) yang lebih interaktif dan premium di storefront, manager, dan cashier."),
            ("2026-07-12", "Meningkatkan keamanan web: menambahkan in-memory rate limiter untuk login/register, menerapkan validasi kekuatan kata sandi (huruf besar/kecil, angka, simbol), dan mengonfigurasi Content Security Policy (CSP) pada Helmet."),
            ("2026-07-12", "Mengoptimasi performa backend: meningkatkan pg pool size ke 20, mengatasi bottleneck kueri N+1 pada produk via bulk fetch addons, mengimplementasi cache matriks otorisasi role, menggabungkan kueri statistik dashboard dengan CTE, mempercepat cold-start migrasi via check to_regclass, serta menambahkan masa berlaku pembayaran Snap 15 menit."),
            ("2026-07-12", "Menambahkan migrasi 0003_outlet_products.sql, kontrak assignment, query harga efektif, endpoint assignment per outlet, filter kategori/katalog storefront, validasi checkout server-side, dan UI pengaturan produk pada outlet aktif."),
            ("2026-07-12", "Membatasi CSS footer storefront ke .app-shell > footer dan menambahkan sistem visual bersama pada manager.css agar seluruh modul Manager/Admin memakai panel, kartu, tabel, toolbar, modal, action area, hover/focus, dan layout responsif yang konsisten."),
            ("2026-07-12", "Menambahkan tabel outlets, foreign key outlet pada user/order/inventory/keuangan, unique SKU per outlet, resolver X-Outlet-Id, CRUD outlet, scoping query, pemilih outlet, dan permission outlets."),
            ("2026-07-12", "Menambahkan tabel payments, adapter Midtrans Snap, webhook SHA-512, simulator lokal, status pembayaran pada order, redirect checkout, dan guard order online sebelum lunas."),
            ("2026-07-12", "Mengganti footer aksi RBAC dengan div.rbac-actions, menambahkan helper copy, state hover/focus pada tabel dan switch, serta toolbar terang yang responsif agar tidak terpengaruh CSS footer global."),
            ("2026-07-12", "Mendesain ulang CashierGrid dan CSS kartu cashier dengan struktur informasi baru, status akses login, footer terang, tombol berlabel, hover, dan focus state."),
            ("2026-07-12", "Mendesain ulang CategoryGrid dan CSS kartu kategori: struktur informasi baru, status visibilitas, footer terang, hover transition, serta layout responsif."),
            ("2026-07-12", "Menyunting bahasa TSD agar lebih natural dan mudah dipahami sambil mempertahankan detail arsitektur, route, API, dan konfigurasi."),
            ("2026-07-12", "Menambahkan helper homePathForRole, memetakan /admin ke dashboard berbasis ManagerApp, memisahkan redirect login Admin/Manager, dan melakukan canonical redirect sesuai role sesi."),
            ("2026-07-10", "Menghapus ketergantungan development pada Netlify Database/Vite plugin, menambahkan API lokal port 3001, proxy Vite /api, PostgreSQL lokal via DATABASE_URL, dan panduan local-only database."),
            ("2026-07-06", "Mengaktifkan production deploy Netlify tokokopdes.netlify.app, provisioning Netlify Database production, apply migration production, setting NETLIFY_DB_URL, dan fallback connection string explicit untuk deploy manual CLI."),
            ("2026-07-06", "Memigrasikan backend ke Express Netlify Function, database ke Netlify PostgreSQL, query async/transactional, migration otomatis, Vite emulator lokal, dan routing SPA/API melalui netlify.toml."),
            ("2026-07-06", "Menambahkan render.yaml untuk Web Service gratis region Singapura, build/start command, health check, secret env, auto-deploy branch main, dan dokumentasi filesystem ephemeral."),
            ("2026-07-05", "Menambahkan role_permissions, default matrix RBAC, endpoint /api/permissions/me dan /api/admin/rbac, guard permission per modul, serta UI tab RBAC Admin."),
            ("2026-07-05", "Menambahkan menu_categories, seed/migrasi kategori, endpoint kategori publik/manager, validasi produk ke kategori database, tab Kategori khusus Manager, dan dropdown produk berbasis API."),
            ("2026-07-05", "Menambahkan financial_entries, payment/diskon order, linkage inventory-produk, pengurangan stok atomik, agregasi Report, endpoint, UI, dan ekspor CSV."),
            ("2026-07-05", "Menambahkan generator dan artefak User Guide Markdown/Word serta menjadikannya dokumen wajib dalam sinkronisasi dokumentasi."),
            ("2026-07-05", "Menambahkan migrasi role admin, seed akun Admin, akses endpoint Manager/Cashier, dan dashboard Admin berbasis ManagerApp."),
            ("2026-07-05", "Menambahkan product_addons, CRUD add-on, pemilih add-on storefront, kalkulasi server-side, dan snapshot add-on pada order_items."),
            ("2026-07-05", "Menambahkan inventory_items, stock_movements, transaksi stok, endpoint inventory, dan UI Inventory untuk Manager/Admin."),
            ("2026-07-05", "Menambahkan updateCashier dan deleteCashier, endpoint PUT/DELETE cashier, form edit/status/password, aksi hapus, serta validasi sesi user aktif."),
            ("2026-07-05", "Memindahkan seed promosi setelah mapper promosi siap agar server tidak gagal saat database masih kosong."),
            ("2026-07-05", "Menambahkan fungsi database getCashiers, endpoint manager cashier, API client, dan tab Cashier di dashboard manager."),
            ("2026-07-05", "Menambahkan app_settings, API settings, hook brand frontend, tab Franchise manager, upload gambar brand, prefix order dinamis, dan key localStorage generik."),
            ("2026-07-05", "Mengganti default database/env/package/README menjadi generik dan membersihkan sisa string brand lama dari kode utama."),
            ("2026-07-05", "Menambahkan output Word untuk FSD dan TSD dengan format dokumen spesifikasi yang lebih rapi."),
        ),
        widths=(1.35, 5.15),
    )

    doc.save(OUTPUTS["tsd"])


def build_user_guide() -> None:
    doc = new_doc(
        "USER GUIDE",
        "Panduan Pengguna",
        "Panduan operasional ringkas untuk Pelanggan, Cashier, Manager, dan Admin pada Franchise Ordering Platform.",
    )

    add_info_box(
        doc,
        "Cara menggunakan panduan",
        "Mulailah dari bagian yang sesuai dengan role Anda. Ikuti langkahnya secara berurutan, lalu gunakan bagian Pemecahan Masalah jika aplikasi tidak berjalan seperti yang diharapkan.",
    )

    add_heading(doc, "1. Mulai Menggunakan Aplikasi", 1)
    add_numbered(
        doc,
        (
            "Buka terminal, lalu masuk ke folder project.",
            "Pastikan service PostgreSQL lokal sudah berjalan.",
            "Jika nama user, password, atau database PostgreSQL berbeda dari konfigurasi bawaan, salin .env.example menjadi .env lalu sesuaikan DATABASE_URL.",
            "Jalankan npm run dev.",
            "Tunggu sampai API lokal dan Vite aktif tanpa error.",
            "Buka http://localhost:5175 melalui browser.",
        ),
    )
    add_info_box(
        doc,
        "Database lokal",
        "Default DATABASE_URL adalah postgres://postgres@127.0.0.1:5432/postgres. Jika PostgreSQL memakai password, gunakan format postgres://postgres:password_anda@127.0.0.1:5432/postgres.",
    )
    add_table(
        doc,
        ("Halaman", "Alamat"),
        (
            ("Toko", "http://localhost:5175/"),
            ("Login", "http://localhost:5175/login"),
            ("Pesanan pelanggan", "http://localhost:5175/orders"),
            ("Stasiun Cashier", "http://localhost:5175/cashier"),
            ("Dashboard Manager", "http://localhost:5175/manager"),
            ("Dashboard Admin", "http://localhost:5175/admin"),
        ),
        widths=(2.05, 4.45),
    )

    add_heading(doc, "1.1 Akun Lokal Bawaan", 2)
    add_table(
        doc,
        ("Role", "Email", "Password"),
        (
            ("Cashier", "cashier@franchise.local", "cashier123"),
            ("Manager", "manager@franchise.local", "manager123"),
            ("Admin", "admin@franchise.local", "admin123"),
        ),
        widths=(1.25, 3.4, 1.85),
    )
    add_info_box(doc, "Penting untuk produksi", "Ganti seluruh password bawaan dan APP_JWT_SECRET sebelum aplikasi dipublikasikan.")

    add_heading(doc, "2. Ringkasan Hak Akses", 1)
    add_body(doc, "Akses ke modul operasional diatur melalui RBAC. Tabel berikut menunjukkan pengaturan awal; Admin dapat menyesuaikannya kapan saja melalui tab RBAC.")
    add_table(
        doc,
        ("Modul", "Pelanggan", "Cashier", "Manager", "Admin"),
        (
            ("Toko dan katalog", "Ya", "Ya", "Ya", "Ya"),
            ("Keranjang dan checkout", "Ya", "-", "-", "-"),
            ("Pesanan Saya", "Ya", "-", "-", "-"),
            ("Stasiun Cashier", "-", "Ya", "Ya", "Ya"),
            ("Kategori menu", "-", "-", "Ya", "Ya"),
            ("Produk dan add-on", "-", "-", "Ya", "Ya"),
            ("Promosi", "-", "-", "Ya", "Ya"),
            ("Akun Cashier", "-", "-", "Ya", "Ya"),
            ("Inventory", "-", "-", "Ya", "Ya"),
            ("Report operasional & keuangan", "-", "-", "Ya", "Ya"),
            ("Outlet", "-", "-", "Ya", "Ya"),
            ("Pengaturan franchise", "-", "-", "Ya", "Ya"),
            ("RBAC modul", "-", "-", "-", "Ya"),
        ),
        widths=(2.1, 1.1, 1.1, 1.1, 1.1),
    )

    add_heading(doc, "3. Panduan Pelanggan", 1)
    add_heading(doc, "3.1 Registrasi dan Login", 2)
    add_numbered(
        doc,
        (
            "Tekan tombol masuk pada halaman toko dan pilih role Pelanggan.",
            "Untuk akun baru, pilih Daftar sekarang lalu isi nama, email, dan password minimal 8 karakter yang mengandung kombinasi huruf besar, huruf kecil, angka, dan simbol khusus.",
            "Untuk akun lama, isi email dan password lalu tekan Masuk.",
        ),
    )
    add_info_box(doc, "Keamanan Akses Akun", "Kata sandi wajib memiliki minimal 8 karakter dan mengandung kombinasi huruf besar, huruf kecil, angka, serta simbol khusus. Sistem membatasi upaya masuk (login/register) maksimal 5 kali percobaan per 15 menit untuk setiap IP.")

    add_heading(doc, "3.2 Memilih Produk dan Add-on", 2)
    add_numbered(
        doc,
        (
            "Cari produk melalui kategori atau kolom pencarian.",
            "Tekan tombol tambah pada kartu produk.",
            "Jika produk memiliki add-on, pilih satu atau beberapa tambahan yang diinginkan.",
            "Periksa total harga per item lalu tekan Tambah ke keranjang.",
        ),
    )
    add_body(doc, "Produk yang sama dengan kombinasi add-on berbeda ditampilkan sebagai baris keranjang yang berbeda.")
    add_info_box(doc, "Katalog mengikuti outlet", "Daftar produk dan harga dapat berbeda antar outlet. Katalog menyesuaikan assignment, status ketersediaan, dan harga khusus outlet yang sedang dipilih.")

    add_heading(doc, "3.3 Keranjang dan Checkout", 2)
    add_numbered(
        doc,
        (
            "Pilih outlet pada bagian atas toko. Pesanan dan stok akan tercatat pada cabang tersebut.",
            "Buka ikon keranjang dan periksa produk, add-on, harga satuan, serta subtotal.",
            "Gunakan tombol tambah, kurang, atau hapus untuk mengatur jumlah.",
            "Tekan Lanjut checkout lalu pilih Diantar atau Ambil sendiri.",
            "Pilih pembayaran Tunai, QRIS, E-wallet, atau Transfer bank serta kode promo jika tersedia.",
            "Isi data pemesan, alamat jika diperlukan, dan catatan opsional.",
            "Untuk tunai, buat pesanan seperti biasa. Untuk QRIS, e-wallet, atau transfer bank, tekan Lanjut ke pembayaran.",
            "Tanpa API key Midtrans, gunakan Simulasikan berhasil/gagal pada halaman lokal. Dengan Midtrans, selesaikan pembayaran pada halaman provider.",
        ),
    )
    add_info_box(doc, "Batas Waktu Pembayaran Online", "Transaksi online (QRIS, E-wallet, Transfer bank) memiliki batas waktu pembayaran 15 menit. Jika tidak dibayar dalam 15 menit, pesanan akan kedaluwarsa secara otomatis dan stok inventaris akan dibebaskan kembali.")
    add_info_box(doc, "Ganti outlet", "Mengganti outlet akan mengosongkan keranjang agar validasi stok tidak tercampur antar cabang.")

    add_heading(doc, "3.4 Melacak Pesanan", 2)
    add_numbered(
        doc,
        (
            "Buka menu profil dan pilih Pesanan saya, atau buka /orders.",
            "Gunakan tab Aktif untuk pesanan berjalan dan Riwayat untuk pesanan selesai/dibatalkan.",
            "Pilih pesanan untuk melihat status, item, add-on, metode penerimaan, alamat, dan total.",
            "Gunakan tombol Perbarui status jika tidak ingin menunggu pembaruan otomatis 15 detik.",
            "Periksa nama outlet dan status pembayaran. Gunakan Bayar sekarang jika pembayaran online masih pending.",
        ),
    )

    add_heading(doc, "4. Panduan Cashier", 1)
    add_heading(doc, "4.1 Login", 2)
    add_numbered(doc, ("Buka halaman login.", "Pilih role Cashier.", "Masukkan akun yang diberikan Manager/Admin.", "Aplikasi membuka /cashier dan mengunci data pada outlet penempatan akun."))
    add_heading(doc, "4.2 Memproses Pesanan", 2)
    add_numbered(
        doc,
        (
            "Gunakan filter atau pencarian untuk menemukan pesanan.",
            "Buka Detail untuk memeriksa item, add-on, pelanggan, alamat, dan catatan.",
            "Ubah Pesanan baru menjadi Sedang dimasak.",
            "Ubah Sedang dimasak menjadi Siap.",
            "Untuk delivery, ubah Siap menjadi Sedang diantar lalu Selesai. Untuk pickup, ubah Siap langsung menjadi Selesai.",
            "Gunakan Dibatalkan hanya jika pesanan benar-benar dibatalkan.",
            "Pesanan online belum dapat diproses sebelum pembayaran berstatus paid.",
        ),
    )

    add_heading(doc, "5. Panduan Manager", 1)
    add_body(doc, "Tab yang terlihat pada dashboard Manager mengikuti permission RBAC dari Admin. Jika sebuah tab tidak muncul, periksa pengaturan akses sebelum menganggap modul tersebut bermasalah.")
    add_body(doc, "Semua tab menggunakan pola tampilan yang sama. Judul dan tombol utama berada di bagian atas panel, isi ditampilkan dalam kartu atau tabel terang, sedangkan tindakan edit, hapus, aktif/nonaktif, dan simpan diletakkan pada area yang mudah ditemukan. Pada ponsel, kartu otomatis tersusun satu kolom dan navigasi tersedia di bagian bawah layar.")
    add_heading(doc, "5.1 Kategori Menu", 2)
    add_numbered(
        doc,
        (
            "Buka tab Kategori lalu tekan Tambah kategori.",
            "Setiap kartu menampilkan ikon, urutan, jumlah produk yang terhubung, serta keterangan apakah kategori sedang tampil di toko atau disembunyikan.",
            "Isi icon/emoji, nama kategori, urutan tampil, dan status aktif.",
            "Gunakan angka urutan lebih kecil agar kategori tampil lebih awal di storefront.",
            "Gunakan switch di bagian bawah kartu untuk menampilkan atau menyembunyikan kategori dari pelanggan.",
            "Gunakan ikon pensil untuk mengubah kategori dan ikon tempat sampah untuk menghapus atau mengarsipkan.",
            "Jika kategori masih dipakai produk, sistem menonaktifkan/mengarsipkan kategori agar produk lama tetap aman.",
            "Jika nama kategori diubah, produk yang memakai kategori tersebut ikut diperbarui otomatis.",
        ),
    )

    add_heading(doc, "5.2 Produk dan Add-on", 2)
    add_numbered(
        doc,
        (
            "Pilih outlet melalui selector Outlet Aktif, lalu buka tab Produk.",
            "Tekan Tambah produk untuk membuat master produk baru; produk otomatis ditugaskan ke outlet aktif.",
            "Isi identitas produk, harga master, kategori, status, dan foto. Pastikan kategori sudah dibuat di tab Kategori.",
            "Pada Pilihan add-on, tekan Tambah add-on lalu isi nama, harga, dan status aktif.",
            "Tekan Simpan produk, lalu gunakan ikon gedung pada kartu untuk membuka pengaturan outlet.",
            "Aktifkan Jual produk ini di outlet, atur status konfigurasi dan ketersediaan, serta isi Harga khusus outlet bila diperlukan.",
            "Kosongkan harga khusus untuk memakai harga master. Gunakan switch Master aktif untuk menonaktifkan produk pada seluruh outlet.",
        ),
    )

    add_heading(doc, "5.3 Promosi", 2)
    add_numbered(
        doc,
        (
            "Buka tab Promosi dan pilih Tambah promosi atau Edit.",
            "Isi judul, deskripsi, kode, tipe dan nilai diskon, minimum belanja, periode, serta status.",
            "Simpan perubahan atau hapus promosi yang tidak digunakan.",
        ),
    )

    add_heading(doc, "5.4 Mengelola Cashier", 2)
    add_numbered(
        doc,
        (
            "Buka tab Cashier dan tekan Tambah cashier.",
            "Setiap kartu menampilkan inisial avatar, nama, email, dan keterangan apakah akun siap digunakan atau aksesnya sedang dinonaktifkan.",
            "Pilih outlet penempatan, lalu isi nama, email, dan password awal minimal 8 karakter yang mengandung kombinasi huruf besar, huruf kecil, angka, dan simbol khusus.",
            "Tekan Edit untuk mengubah outlet, identitas, password, atau status login.",
            "Tekan Hapus untuk menghapus akun. Akun nonaktif atau terhapus tidak dapat memakai sesi lama.",
        ),
    )

    add_heading(doc, "5.5 Inventory", 2)
    add_numbered(
        doc,
        (
            "Pilih outlet aktif, buka tab Inventory, lalu tekan Tambah item.",
            "Isi nama, SKU unik, satuan, stok awal, minimum stok, harga modal, dan status.",
            "Pilih Produk terkait dan jumlah Pemakaian per produk terjual jika stok harus berkurang otomatis.",
            "Tekan tombol tambah pada baris item untuk mencatat pergerakan stok.",
            "Pilih Stok masuk, Stok keluar, Koreksi tambah, atau Koreksi kurang; isi jumlah dan catatan.",
            "Buka Riwayat mutasi untuk melihat stok sebelum/sesudah, waktu, catatan, dan pembuat.",
        ),
    )
    add_info_box(doc, "Kontrol stok", "Stok keluar dan koreksi kurang ditolak jika membuat stok negatif. Stok yang sama dengan atau di bawah minimum ditandai Stok menipis.")

    add_heading(doc, "5.6 Report Operasional dan Keuangan", 2)
    add_numbered(
        doc,
        (
            "Pilih outlet aktif, buka tab Report, lalu pilih periode Dari/Sampai; data diperbarui otomatis setiap 30 detik.",
            "Pada Operasional, tinjau penjualan harian, produk terlaris, pembayaran, stok, dan pelanggan.",
            "Pada Keuangan, tinjau laporan laba rugi (menampilkan omzet penjualan, pengeluaran bahan baku / HPP otomatis, laba kotor, biaya operasional manual, dan laba bersih), arus kas, neraca, perubahan modal, dan biaya per kategori.",
            "Tekan Catat transaksi untuk memasukkan biaya operasional (seperti sewa, listrik, gaji, dll. - hindari mencatat pembelian bahan baku fisik di sini karena biaya bahan baku sudah dihitung otomatis oleh sistem sebagai HPP saat produk terjual), modal masuk, atau penarikan modal.",
            "Tekan Ekspor CSV untuk mengunduh seluruh bagian laporan pada periode terpilih (termasuk HPP dan Laba Kotor).",
        ),
    )
    add_info_box(doc, "Batas laporan", "Penjualan non-batal masuk otomatis dan bahan baku yang digunakan otomatis dihitung sebagai HPP. Pembatalan pesanan otomatis mengembalikan stok ke inventory. Laporan bersifat manajerial dasar dan belum menghitung utang, piutang, depresiasi, atau pajak.")

    add_heading(doc, "5.7 Mengelola Outlet", 2)
    add_numbered(doc, (
        "Buka tab Outlet lalu tekan Tambah outlet.",
        "Isi kode unik, nama cabang, alamat, telepon, status aktif, dan opsi outlet utama.",
        "Gunakan Pilih outlet atau selector kanan atas untuk mengganti konteks Cashier, Inventory, dan Report.",
        "Outlet utama tidak dapat dihapus. Outlet dengan histori akan diarsipkan.",
    ))
    add_info_box(doc, "Data bersama dan terpisah", "Kategori, master produk, add-on, promosi, dan brand digunakan bersama. Produk yang tampil serta harganya mengikuti assignment outlet. Pesanan, cashier, inventory, keuangan, dan laporan dipisahkan per outlet.")

    add_heading(doc, "5.8 Pengaturan Franchise", 2)
    add_numbered(
        doc,
        (
            "Buka tab Franchise.",
            "Ubah identitas brand, warna, kontak, prefix pesanan, konten halaman depan, dan gambar.",
            "Tekan Simpan pengaturan franchise lalu periksa halaman toko.",
        ),
    )

    add_heading(doc, "6. Panduan Admin", 1)
    add_numbered(
        doc,
        (
            "Pilih role Admin pada halaman login.",
            "Setelah login berhasil, aplikasi membuka /admin dan menampilkan label Admin Dashboard.",
            "Gunakan modul yang diizinkan RBAC; secara default Kategori, Produk, Promosi, Cashier, Inventory, Report, Outlet, Franchise, Stasiun Cashier, dan RBAC aktif.",
            "Buka tab RBAC untuk mengatur modul yang boleh diakses Cashier, Manager, dan Admin.",
            "Centang modul yang ingin diaktifkan atau hilangkan centang untuk menonaktifkan. Periksa kembali pengaturannya pada panel Simpan pengaturan akses, lalu tekan Simpan RBAC. Gunakan Reset tampilan untuk membatalkan perubahan yang belum disimpan.",
            "Permission Admin - RBAC wajib aktif dan tidak dapat dimatikan agar Admin tidak terkunci dari pengaturan akses.",
            "Jika akses modul dimatikan, tab modul tersebut hilang dari dashboard role terkait dan request backend ditolak.",
            "Buka Stasiun cashier untuk melihat dan memproses semua pesanan jika permission Stasiun Cashier aktif.",
        ),
    )
    add_info_box(doc, "URL sesuai role", "Alamat dashboard selalu mengikuti role yang sedang login. Admin yang membuka /manager akan dibawa kembali ke /admin, sedangkan Manager yang membuka /admin akan diarahkan ke /manager.")

    add_heading(doc, "7. Menu Profil", 1)
    add_body(doc, "Menu profil tersedia di pojok kanan atas untuk setiap role. Dari menu ini, pengguna dapat:")
    add_bullets(
        doc,
        (
            "Membuka halaman utama sesuai role.",
            "Mengubah nama dan email.",
            "Mengganti password.",
            "Keluar dari aplikasi.",
        ),
    )

    add_heading(doc, "8. Referensi Status Pesanan", 1)
    add_table(
        doc,
        ("Status", "Arti"),
        (
            ("Pesanan baru", "Pesanan diterima dan menunggu diproses."),
            ("Sedang dimasak", "Tim sedang menyiapkan pesanan."),
            ("Siap", "Pesanan siap diambil atau dikirim."),
            ("Sedang diantar", "Pesanan sedang menuju pelanggan."),
            ("Selesai", "Pesanan telah diterima/diambil."),
            ("Dibatalkan", "Pesanan tidak dilanjutkan."),
        ),
        widths=(1.9, 4.6),
    )

    add_heading(doc, "9. Pemecahan Masalah", 1)
    add_heading(doc, "9.1 Halaman Tidak Dapat Dibuka", 2)
    add_bullets(doc, ("Pastikan terminal berada di folder project.", "Pastikan PostgreSQL lokal menyala.", "Jalankan npm run dev dan biarkan terminal tetap terbuka.", "Jika port 5175 dipakai, tutup proses lama sebelum menjalankan ulang project.", "Jika port 3001 dipakai aplikasi lain, tutup proses lama atau ubah PORT untuk API lokal."))
    add_heading(doc, "9.2 Database Lokal Gagal Tersambung", 2)
    add_bullets(doc, ("Pastikan service PostgreSQL berjalan di komputer.", "Pastikan DATABASE_URL di .env sesuai user, password, host, port, dan nama database.", "Jika belum membuat database khusus, gunakan database bawaan postgres.", "Development lokal tidak memakai database Netlify; error koneksi harus diperbaiki di PostgreSQL lokal."))
    add_heading(doc, "9.3 Sesi Berakhir atau Akses Ditolak", 2)
    add_bullets(doc, ("Logout lalu login kembali menggunakan role yang benar.", "Pastikan akun masih aktif.", "Hubungi Manager/Admin jika akun Cashier dinonaktifkan atau permission modul belum diberikan."))
    add_heading(doc, "9.4 Add-on Gagal Digunakan", 2)
    add_body(doc, "Add-on mungkin dinonaktifkan setelah masuk keranjang. Hapus produk dari keranjang lalu tambahkan kembali dengan add-on yang tersedia.")
    add_heading(doc, "9.5 Kategori Tidak Muncul di Storefront", 2)
    add_bullets(doc, ("Pastikan role Anda memiliki permission Kategori di RBAC.", "Pastikan kategori berstatus aktif.", "Pastikan produk terkait aktif, sudah ditugaskan, dan tersedia pada outlet yang dipilih.", "Kategori hanya tampil jika memiliki produk yang dapat dijual pada outlet aktif."))
    add_heading(doc, "9.6 Mutasi Stok Ditolak", 2)
    add_bullets(doc, ("Pastikan jumlah lebih dari nol.", "Pastikan pengurangan tidak membuat stok negatif.", "Pastikan item inventory masih aktif."))
    add_heading(doc, "9.7 Upload Gambar Gagal", 2)
    add_bullets(doc, ("Gunakan PNG, JPG/JPEG, WebP, atau GIF.", "Gunakan file maksimal sekitar 2 MB."))
    add_heading(doc, "9.8 Pembayaran Online Tidak Terbuka", 2)
    add_bullets(doc, ("Pastikan outlet aktif dipilih dan pelanggan sudah login.", "Isi MIDTRANS_SERVER_KEY, MIDTRANS_CLIENT_KEY, MIDTRANS_IS_PRODUCTION, dan PUBLIC_APP_URL lalu restart server.", "Atur Payment Notification URL ke https://domain-anda/api/payments/midtrans/notification.", "Jika Server Key kosong, aplikasi memakai /payment-simulator.", "Pembayaran gagal/kedaluwarsa membatalkan order dan mengembalikan stok."))
    add_heading(doc, "9.9 Data Outlet Terlihat Kosong", 2)
    add_bullets(doc, ("Periksa outlet aktif pada selector dashboard.", "Untuk produk, buka ikon gedung dan pastikan assignment, konfigurasi aktif, serta Produk tersedia dijual sudah aktif.", "Cashier hanya melihat outlet penempatannya.", "Pesanan, inventory, keuangan, dan laporan memang dipisahkan per outlet."))
    add_heading(doc, "9.10 Report Tidak Menampilkan Transaksi", 2)
    add_bullets(doc, ("Periksa periode Dari dan Sampai.", "Order cancelled tidak dihitung.", "Pastikan login menggunakan role Manager atau Admin."))
    add_heading(doc, "9.11 Mengakses Versi Hosting Publik", 2)
    add_bullets(
        doc,
        (
            "Buka alamat publik https://tokokopdes.netlify.app; pengunjung tidak memerlukan verifikasi tambahan.",
            "Website publik tidak dapat memakai database lokal di laptop jika laptop mati/offline.",
            "Jika ingin hosting publik aktif 24 jam, gunakan PostgreSQL online dan isi DATABASE_URL/POSTGRES_URL pada environment hosting.",
            "Jika project tidak dapat diakses dan limit Free habis, tunggu reset periode atau tingkatkan paket dari dashboard Netlify.",
            "Kredensial publik mengikuti password environment variable yang disimpan di hosting, bukan password yang disimpan di Git.",
        ),
    )

    add_heading(doc, "10. Aturan Pembaruan Dokumen", 1)
    add_info_box(
        doc,
        "Wajib diperbarui bersama",
        "Setiap perubahan fitur, role, API, database, alur pengguna, atau UI utama wajib memperbarui FSD, TSD, dan User Guide dalam format Markdown dan Word.",
    )
    add_bullets(doc, ("docs/FSD.md dan docs/FSD.docx", "docs/TSD.md dan docs/TSD.docx", "docs/USER_GUIDE.md dan docs/USER_GUIDE.docx"))

    add_heading(doc, "11. Riwayat Perubahan", 1)
    add_table(
        doc,
        ("Tanggal", "Perubahan"),
        (
            ("2026-07-12", "Mengganti pemilih outlet bawaan browser dengan komponen dropdown kustom (OutletSelector) yang lebih interaktif dan premium di storefront, manager, dan cashier."),
            ("2026-07-12", "Meningkatkan keamanan web: menambahkan in-memory rate limiter untuk login/register, menerapkan validasi kekuatan kata sandi (huruf besar/kecil, angka, simbol), dan mengonfigurasi Content Security Policy (CSP) pada Helmet."),
            ("2026-07-12", "Mengoptimasi performa backend: meningkatkan pg pool size ke 20, mengatasi bottleneck kueri N+1 pada produk via bulk fetch addons, mengimplementasi cache matriks otorisasi role, menggabungkan kueri statistik dashboard dengan CTE, mempercepat cold-start migrasi via check to_regclass, serta menambahkan masa berlaku pembayaran Snap 15 menit."),
            ("2026-07-12", "Menambahkan panduan master produk global dan assignment per outlet, termasuk harga khusus, status aktif/tersedia, dampak pada katalog, serta pemecahan masalah produk outlet."),
            ("2026-07-12", "Menambahkan panduan tampilan seragam seluruh modul Manager/Admin, termasuk letak tindakan, pola panel terang, dan penyesuaian navigasi pada layar ponsel."),
            ("2026-07-12", "Menambahkan panduan memilih dan mengelola outlet, menempatkan cashier, serta memahami pemisahan pesanan, stok, dan laporan per cabang."),
            ("2026-07-12", "Menambahkan panduan pembayaran online Midtrans, simulator lokal, status pembayaran pada tracking, webhook, dan pemecahan masalah konfigurasi."),
            ("2026-07-12", "Memperbarui panduan RBAC sesuai tampilan baru: tabel lebih nyaman dibaca serta panel Reset dan Simpan yang ringkas di bawah daftar akses."),
            ("2026-07-12", "Memperbarui panduan kartu cashier agar sesuai dengan tampilan baru, termasuk status akses dan tombol aksi berlabel."),
            ("2026-07-12", "Memperbarui panduan kartu kategori agar sesuai dengan tampilan baru, termasuk informasi visibilitas, switch status, dan tombol aksi."),
            ("2026-07-12", "Merapikan bahasa User Guide agar terasa lebih natural dan instruksinya lebih mudah diikuti tanpa mengubah langkah operasional."),
            ("2026-07-12", "Memisahkan alamat dashboard Admin menjadi /admin dan Manager menjadi /manager, termasuk pengalihan otomatis jika URL tidak sesuai role login."),
            ("2026-07-10", "Menambahkan panduan menjalankan database PostgreSQL lokal tanpa Netlify, API lokal port 3001, Vite port 5175, dan catatan bahwa hosting publik membutuhkan database online."),
            ("2026-07-06", "Menambahkan URL live Netlify https://tokokopdes.netlify.app, catatan deploy context production, dan database production yang sudah diprovision."),
            ("2026-07-06", "Menambahkan panduan akses Netlify, API satu domain, PostgreSQL persisten, credential environment, dan batas kredit paket Free."),
            ("2026-07-06", "Menambahkan panduan akses deployment Render, waktu bangun service gratis, credential deploy, dan batasan data SQLite pada filesystem sementara."),
            ("2026-07-05", "Menambahkan panduan RBAC Admin per modul, default hak akses, efek tab dashboard berbasis permission, dan pemecahan masalah akses ditolak."),
            ("2026-07-05", "Menambahkan panduan kategori menu custom khusus Manager, efek kategori aktif/nonaktif pada storefront, dan catatan akses Admin terhadap kategori."),
            ("2026-07-05", "Menambahkan panduan Report, pembayaran/promo, transaksi biaya/modal, ekspor CSV, harga modal, dan stok otomatis."),
            ("2026-07-05", "Membuat User Guide pertama untuk seluruh role, add-on produk, Inventory, profil, dan pemecahan masalah."),
        ),
        widths=(1.35, 5.15),
    )

    doc.save(OUTPUTS["user_guide"])


def main() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    build_fsd()
    build_tsd()
    build_user_guide()
    for name, path in OUTPUTS.items():
        print(f"{name.upper()} -> {path}")


if __name__ == "__main__":
    main()
